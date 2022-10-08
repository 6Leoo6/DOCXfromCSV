from docx import Document
from io import BytesIO
from zipfile import ZipFile


async def convertCSV(csv, model):
    # Read the bytes of the files uploaded
    csv_bytes = await csv.read()
    model_bytes = BytesIO(await model.read())

    # Parse the csv
    csv = csv_bytes.decode('utf-8').replace('\r', '').split('\n')
    csv = list(map(lambda x: x.split(','), csv))

    # Separate the first row from the others
    headers = csv.pop(0)
    rows = csv

    # Load the word doc
    doc = Document(model_bytes)
    # Create an empty zip file
    zip_bytes = BytesIO()
    # Write into the zip variable
    with ZipFile(zip_bytes, 'w') as zip:
        # Loop through the rows
        for csv_row in rows:
            full_doc = []
            #Set the filenamne to the first value of the row
            filename = csv_row[0]

            # Loop through each paragraph found in the model
            for i, paragraph in enumerate(doc.paragraphs):
                # Parse the paragraph text
                paragraph = paragraph.text
                
                #Loop through each column title found in the csv
                for col, header in enumerate(headers):
                    #Replace the title with the corresponding value
                    paragraph = paragraph.replace(
                        '${'+header+'}', csv_row[col])
                #If it's the first line of the model and it starts with [ and ends with ] then set the filename to it
                if i == 0 and paragraph.startswith('[') and paragraph.endswith(']'):
                    #Remove the square brackets from both ends
                    filename = paragraph.strip('[]')
                    continue
                #If it's not the filename then add it to the list of paragraphs
                full_doc.append(paragraph)

            #Create an empty document
            generated_doc = Document()
            #Loop through the parsed paragraphs
            for paragraph in full_doc:
                #Add back the paragraphs
                generated_doc.add_paragraph(paragraph)
            #Save the doc into memory
            doc_bytes = BytesIO()
            generated_doc.save(doc_bytes)
            #Add the file to the zip archive
            zip.writestr(f'{filename}.docx', doc_bytes.getvalue())

    #Return the bytes of the zip
    return zip_bytes.getvalue()